#!/usr/bin/env python3
"""
Pre-processing script for the CS50 Edudata dataset.
This script prepares the raw Edudata folder so it matches the structure
expected by the llm pipelines (transcript, notes, qna, ppt).

It does the following:
1. Re-organizes files: Maps the root .webm and .srt files to their respective lesson folders (0, 1, 2, ...).
2. Converts .srt files to plain text transcripts (transcript.txt).
3. Parses index.html files in lesson folders and converts them to plain text notes (notes.txt).
4. Extracts frames from the .webm video files (1 frame every 60 seconds) and saves them as images,
   and compiles them into a `ppt.pdf` file to be compatible with pipelines.

Prerequisites:
- pip install beautifulsoup4 opencv-python pillow reportlab
"""

import os
import re
import cv2
from PIL import Image
from bs4 import BeautifulSoup
from reportlab.pdfgen import canvas
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def clean_srt(srt_file_path, docx_file_path):
    """Convert .srt to DOCX."""
    if not os.path.exists(srt_file_path):
        return
    logger.info(f"Converting SRT: {srt_file_path}")
    with open(srt_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    transcript = []
    for line in lines:
        line = line.strip()
        if not line or line.isdigit() or '-->' in line:
            continue
        transcript.append(line)
        
    from docx import Document
    doc = Document()
    full_text = ' '.join(transcript).replace('  ', ' ')
    doc.add_paragraph(full_text)
    doc.save(docx_file_path)
    logger.info(f"Saved: {docx_file_path}")

def clean_html_notes(html_file_path, docx_file_path):
    """Convert HTML document into DOCX notes."""
    if not os.path.exists(html_file_path):
        return
    logger.info(f"Converting HTML: {html_file_path}")
    with open(html_file_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')
        
    text = soup.get_text(separator='\n', strip=True)
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    doc.save(docx_file_path)
    logger.info(f"Saved: {docx_file_path}")

def combine_html_in_dir(dir_path, output_docx_path):
    """Combine all HTML files in directory and subdirectories into one notes.docx"""
    from docx import Document
    doc = Document()
    added = False
    
    for root, dirs, files in os.walk(dir_path):
        for file in files:
            if file.endswith('.html'):
                html_path = os.path.join(root, file)
                with open(html_path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f, 'html.parser')
                    rel_path = os.path.relpath(html_path, dir_path)
                    doc.add_heading(f"Section: {rel_path}", level=1)
                    doc.add_paragraph(soup.get_text(separator='\n', strip=True))
                    added = True
                    
    if added:
        doc.save(output_docx_path)
        logger.info(f"Saved compiled notes: {output_docx_path}")

def extract_frames_to_pdf(video_path, output_pdf_path, frame_rate=1/60.0):
    """
    Extracts 1 frame every X seconds (default 60) and compiles them into a single PDF ('ppt.pdf').
    """
    if not os.path.exists(video_path):
        return
    logger.info(f"Extracting frames from video: {video_path} -> {output_pdf_path}")
    
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        logger.error(f"Failed to open video: {video_path}")
        return

    fps = cap.get(cv2.CAP_PROP_FPS)
    if not fps or fps != fps:
        fps = 30.0 # fallback

    frame_interval = int(fps / frame_rate) # e.g. 30 fps * 60 seconds = 1800 frames
    
    frames_extracted = 0
    images = []
    
    success, frame = cap.read()
    count = 0
    
    while success:
        if count % frame_interval == 0:
            # Convert BGR (cv2) to RGB
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            img = Image.fromarray(frame_rgb)
            images.append(img)
            frames_extracted += 1
            if frames_extracted % 10 == 0:
                logger.info(f"  ... extracted {frames_extracted} frames")
                
        success, frame = cap.read()
        count += 1

    cap.release()
    
    if images:
        logger.info(f"Saving {len(images)} frames to {output_pdf_path}")
        images[0].save(
            output_pdf_path, "PDF" ,resolution=100.0, save_all=True, append_images=images[1:]
        )
    else:
        logger.warning(f"No frames extracted for {video_path}")

def main():
    # Adjust path assuming this script is in CS50_pipelines
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    # For remote, data might be just 'edudata' inside data folder
    # We will search for 'data/edudata' or 'data/Edudata'
    edudata_dir_candidates = [
        os.path.join(base_dir, "data", "edudata"),
        os.path.join(base_dir, "data", "Edudata"),
        os.path.join(base_dir, "edudata"),
        os.path.join(base_dir, "Edudata")
    ]
    
    edudata_dir = None
    for candidate in edudata_dir_candidates:
        if os.path.exists(candidate):
            edudata_dir = candidate
            break
            
    if not edudata_dir:
        logger.error("Could not find edudata directory!")
        return
        
    logger.info(f"Starting pre-processing on: {edudata_dir}")
    
    # Map lecture numbers to video and srt files
    files_in_root = [f for f in os.listdir(edudata_dir) if os.path.isfile(os.path.join(edudata_dir, f))]
    
    lecture_map = {} # '0' -> {'webm': 'path', 'srt': 'path'}
    
    for f in files_in_root:
        # Expected format: "1 - Introduction - ...webm" or "1 - Intro...srt"
        match = re.match(r'^(\d+)\s*-', f)
        if match:
            lec_num = match.group(1)
            if lec_num not in lecture_map:
                lecture_map[lec_num] = {}
                
            if f.endswith('.webm'):
                lecture_map[lec_num]['webm'] = os.path.join(edudata_dir, f)
            elif f.endswith('.srt'):
                lecture_map[lec_num]['srt'] = os.path.join(edudata_dir, f)
                
    # Now process each numbered directory
    dirs_in_root = [d for d in os.listdir(edudata_dir) if os.path.isdir(os.path.join(edudata_dir, d)) and d.isdigit()]
    
    for d in dirs_in_root:
        target_dir = os.path.join(edudata_dir, d)
        logger.info(f"\n======================\nProcessing Lecture {d}\n======================")
        
        # 1. HTML -> Notes.docx
        notes_docx = os.path.join(target_dir, "notes.docx")
        combine_html_in_dir(target_dir, notes_docx)
        
        # 2. SRT -> Transcript.docx
        if d in lecture_map and 'srt' in lecture_map[d]:
            transcript_docx = os.path.join(target_dir, "transcript.docx")
            clean_srt(lecture_map[d]['srt'], transcript_docx)
            
        # 3. WebM -> PPT.pdf
        if d in lecture_map and 'webm' in lecture_map[d]:
            ppt_pdf = os.path.join(target_dir, "ppt.pdf")
            if not os.path.exists(ppt_pdf):
                extract_frames_to_pdf(lecture_map[d]['webm'], ppt_pdf, frame_rate=1/60.0) # 1 frame per min
            else:
                logger.info(f"PPT already exists: {ppt_pdf}. Skipping video extraction.")

if __name__ == "__main__":
    main()
